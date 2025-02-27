import streamlit as st
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from werkzeug.utils import secure_filename
import tempfile
import shutil
import re
from datetime import datetime

# Set page config first - must be the first Streamlit command
st.set_page_config(
    page_title="Event Report Generator",
    layout="wide"
)

TEMPLATE_PATH = 'workshop_template.docx'  # Path to your report template

def add_underline(run):
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    run._element.rPr.append(u)

def add_centered_image(doc, image_path, width=Cm(9), height=Cm(16)):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=width, height=height)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def ensure_heading_style(doc, style_name, font_name='DIN Pro Regular', font_size=12, rgb_color=(135, 206, 235)):
    styles = doc.styles
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, 1)
    
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(*rgb_color)
    return style

def save_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return None
    
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, secure_filename(uploaded_file.name))
    
    with open(temp_path, 'wb') as f:
        f.write(uploaded_file.getvalue())
    
    return temp_path

def add_signature_section(doc, faculty_name, hod_name):
    # Add space before signatures
    doc.add_paragraph('\n\n')
    
    # Create a table for signatures with specific widths
    signature_table = doc.add_table(rows=2, cols=2)
    signature_table.autofit = False
    
    # Set column widths (50% each)
    for cell in signature_table.columns[0].cells:
        cell.width = Cm(8)
    for cell in signature_table.columns[1].cells:
        cell.width = Cm(8)
    
    # Faculty-in-charge section
    faculty_label_cell = signature_table.cell(0, 0)
    faculty_label_paragraph = faculty_label_cell.paragraphs[0]
    faculty_label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    faculty_label_run = faculty_label_paragraph.add_run("Name & Signature of the")
    faculty_label_run.font.name = 'DIN Pro Regular'
    faculty_label_run.font.size = Pt(11)
    
    faculty_label_cell2 = signature_table.cell(1, 0)
    faculty_label_paragraph2 = faculty_label_cell2.paragraphs[0]
    faculty_label_paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    faculty_label_run2 = faculty_label_paragraph2.add_run("Faculty-in-charge")
    faculty_label_run2.font.name = 'DIN Pro Regular'
    faculty_label_run2.font.size = Pt(11)
    faculty_name_run = faculty_label_paragraph2.add_run(f"\n{faculty_name}")
    faculty_name_run.font.name = 'DIN Pro Regular'
    faculty_name_run.font.size = Pt(11)
    
    # HoD section
    hod_label_cell = signature_table.cell(0, 1)
    hod_label_paragraph = hod_label_cell.paragraphs[0]
    hod_label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    hod_label_run = hod_label_paragraph.add_run("Name & Signature of")
    hod_label_run.font.name = 'DIN Pro Regular'
    hod_label_run.font.size = Pt(11)
    
    hod_label_cell2 = signature_table.cell(1, 1)
    hod_label_paragraph2 = hod_label_cell2.paragraphs[0]
    hod_label_paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    hod_label_run2 = hod_label_paragraph2.add_run("HoD")
    hod_label_run2.font.name = 'DIN Pro Regular'
    hod_label_run2.font.size = Pt(11)
    hod_name_run = hod_label_paragraph2.add_run(f"\n{hod_name}")
    hod_name_run.font.name = 'DIN Pro Regular'
    hod_name_run.font.size = Pt(11)

def process_text_with_bullets(text):
    """Process text to preserve bullet points in docx format"""
    lines = text.split('\n')
    processed_lines = []
    for line in lines:
        # Check if line starts with bullet point (-, *, •)
        if line.strip().startswith(('-', '*', '•')):
            processed_line = line.strip()[1:].strip()  # Remove the bullet and trim
            processed_lines.append((True, processed_line))  # Mark as bullet point
        else:
            processed_lines.append((False, line))  # Regular line
    return processed_lines

def add_text_with_bullets(doc, text):
    """Add text to document with bullet points preserved"""
    processed_lines = process_text_with_bullets(text)
    
    for is_bullet, line in processed_lines:
        if is_bullet:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Pt(36)
            run = p.add_run(line)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
        
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
    
    return doc

def create_report(inputs, files):
    temp_dir = tempfile.mkdtemp()
    try:
        doc = Document(TEMPLATE_PATH)
        
        # Set up styles
        ensure_heading_style(doc, 'Heading 1')
        ensure_heading_style(doc, 'Heading 2')
        
        # Add department name
        header = doc.add_paragraph()
        run = header.add_run(f"Department of {inputs['department_name']}")
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(79, 129, 189)
        add_underline(run)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add event type as title
        title = doc.add_paragraph()
        run = title.add_run(f"{inputs['event_type']}")
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(79, 129, 189)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('\n')
        
        # Create details table
        table = doc.add_table(rows=7, cols=2)  # Ensure 7 rows
        table.style = 'Table Grid'
        
        # Dynamic table content based on event type
        rows = [
            ("Topic", inputs['topic']),
            ("Expert", inputs['expert']) if inputs['event_type'] != 'Field Visit' else None,
            ("Venue", inputs['venue']),
            ("Date", inputs['date']),
            ("Start Time", inputs['start_time']),
            ("Faculty Coordinator", inputs['coordinator']),
            ("Number of Participants", inputs['num_participants'])
        ]
        
        # Filter out None values
        rows = [row for row in rows if row is not None]
        
        # Ensure we have the right number of rows
        while len(rows) < 7:
            # Add empty rows if needed
            rows.append(("", ""))
        
        for i, (label, value) in enumerate(rows[:7]):  # Limit to first 7 rows
            try:
                row = table.rows[i]
                label_cell = row.cells[0]
                label_paragraph = label_cell.paragraphs[0]
                label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                label_run = label_paragraph.add_run(label)
                label_run.font.name = 'DIN Pro Regular'
                label_run.font.size = Pt(11)
                label_run.bold = True

                value_cell = row.cells[1]
                value_paragraph = value_cell.paragraphs[0]
                value_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                value_run = value_paragraph.add_run(str(value))
                value_run.font.name = 'DIN Pro Regular'
                value_run.font.size = Pt(11)
            except IndexError:
                st.error(f"Error in processing table row: {label}")
                continue  # Skip this row if there's an error
        
        # Add sections
        sections = [
            ("Summary of the Event:", inputs['summary']),
            ("Outcome of the Event:", inputs['outcome'])
        ]
        
        for heading_text, content in sections:
            doc.add_paragraph('\n')
            heading = doc.add_paragraph()
            run = heading.add_run(heading_text)
            run.font.name = 'DIN Pro Regular'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(79, 129, 189)
            
            # Add content with bullet points preserved
            add_text_with_bullets(doc, content)
        
        # Add images sections (only if files are uploaded)
        if files['invite_image']:
            doc.add_page_break()
            invite_heading = doc.add_paragraph()
            run = invite_heading.add_run("Invite")
            run.font.name = 'DIN Pro Regular'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(79, 129, 189)
            invite_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            invite_path = save_uploaded_file(files['invite_image'])
            if invite_path:
                add_centered_image(doc, invite_path, width=Cm(9), height=Cm(17))
        
        image_sections = [
            ("Action Photos", files['action_photos'], Cm(10), Cm(10)),
            ("Attendance Sheet", files['attendance_photos'], Cm(9), Cm(9)),
            ("Analysis Report", files['analysis_photos'], Cm(9), Cm(9))
        ]
        
        for heading_text, photos, width, height in image_sections:
            if photos:
                doc.add_page_break()
                heading = doc.add_paragraph()
                run = heading.add_run(heading_text)
                run.font.name = 'DIN Pro Regular'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(79, 129, 189)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph('\n')
                
                for photo in photos:
                    photo_path = save_uploaded_file(photo)
                    if photo_path:
                        add_centered_image(doc, photo_path, width=width, height=height)
        
        # Ensure the signature section is placed at the end of the document, without creating a new page
        add_signature_section(doc, inputs['faculty_in_charge'], inputs['hod_name'])
        
        # Save document with dynamic filename
        event_type = inputs['event_type'].lower().replace(" ", "_")
        file_name = f"{event_type}_1_{inputs['department_name']}.docx"
        output_path = os.path.join(temp_dir, file_name)
        doc.save(output_path)
        
        with open(output_path, 'rb') as f:
            file_data = f.read()
        
        return file_data
        
    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            st.warning(f"Could not clean up temporary files: {str(e)}")

# Function to validate time input in both 12-hour and 24-hour formats
def validate_time(time_input):
    # Check for 24-hour format (HH:MM)
    hour24_pattern = re.compile(r'^([0-1]?[0-9]|2[0-3]):([0-5][0-9])$')
    
    # Check for 12-hour format (HH:MM AM/PM)
    hour12_pattern = re.compile(r'^(1[0-2]|0?[1-9]):([0-5][0-9])\s*(AM|PM|am|pm)$')
    
    if hour24_pattern.match(time_input) or hour12_pattern.match(time_input):
        return True
    return False

# Function to convert time to 24-hour format for consistency
def format_time_to_24hr(time_input):
    # If already in 24-hour format
    hour24_pattern = re.compile(r'^([0-1]?[0-9]|2[0-3]):([0-5][0-9])$')
    if hour24_pattern.match(time_input):
        return time_input
    
    # Try to parse 12-hour format
    try:
        # Parse 12-hour format with AM/PM
        time_obj = datetime.strptime(time_input, '%I:%M %p')
        return time_obj.strftime('%H:%M')
    except ValueError:
        try:
            # Try alternative format (no space between time and AM/PM)
            time_obj = datetime.strptime(time_input, '%I:%M%p')
            return time_obj.strftime('%H:%M')
        except ValueError:
            # If parsing fails, return original input
            return time_input

# Generate time options for dropdown (12-hour format with 15-minute intervals)
def generate_time_options():
    times = []
    for hour in range(6, 22):  # 6 AM to 10 PM
        for minute in [0, 15, 30, 45]:  # 15-minute intervals
            time_str = f"{hour % 12 if hour % 12 != 0 else 12}:{minute:02d} {'AM' if hour < 12 else 'PM'}"
            times.append(time_str)
    return times  # No need to sort as they are already in order

def main():
    st.title("Event Report Generator")
    
    # Dropdown to select event type
    event_type = st.selectbox("Select Event Type", ["Workshop", "Field Visit", "Masterclass"])
    
    # Input fields
    department_name = st.text_input("Department Name")
    topic = st.text_input("Topic")
    expert = st.text_input("Expert Name") if event_type != 'Field Visit' else None
    venue = st.text_input("Venue")
    date = st.date_input("Date")
    
    # Time input with both manual entry and dropdown selection options
    st.write("Start Time")
    time_col1, time_col2 = st.columns([1, 1])
    
    with time_col1:
        # Dropdown selection for time (12-hour format with 15-minute intervals)
        time_options = generate_time_options()
        selected_time = st.selectbox("Select Time", time_options, index=32)  # Default to 9:00 AM
    
    with time_col2:
        # Manual time input option with a proper label
        manual_time = st.text_input("Or enter time manually (HH:MM AM/PM or 24-hour format)", 
                                    placeholder="e.g. 9:30 AM or 09:30",
                                    key="manual_time_input")
    
    # Use the manually entered time if provided, otherwise use the dropdown selection
    start_time = manual_time if manual_time else selected_time
    
    # Validate and format the time input
    if start_time:
        if not validate_time(start_time):
            st.error("Please enter time in HH:MM format (e.g. 09:30) or HH:MM AM/PM format (e.g. 9:30 AM)")
        else:
            # Convert to 24-hour format for consistency in the report
            start_time = format_time_to_24hr(start_time)
    
    # Faculty coordinator with auto-updating Faculty-in-charge
    coordinator = st.text_input("Faculty Coordinator")
    
    # Auto-fill Faculty-in-charge with same value as coordinator
    faculty_in_charge = st.text_input("Name of the Faculty-in-charge", value=coordinator)
    
    num_participants = st.number_input("Number of Participants", min_value=1)
    
    # Larger text areas for summary and outcome with proper labels
    st.write("Summary of the Event")
    summary = st.text_area("Summary of the Event", height=100, key="summary_input", label_visibility="collapsed")
    
    st.write("Outcome of the Event")
    outcome = st.text_area("Outcome of the Event", height=100, key="outcome_input", label_visibility="collapsed")
    
    hod_name = st.text_input("Name of the HoD")
    
    # File uploads
    st.subheader("Upload Images")
    invite_image = st.file_uploader("Upload Invite Image", type=['png', 'jpg', 'jpeg'])
    action_photos = st.file_uploader("Upload Action Photos", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
    attendance_photos = st.file_uploader("Upload Attendance Sheet Photos", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
    analysis_photos = st.file_uploader("Upload Analysis Report Photos", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
    
    if st.button("Generate Report"):
        # Validate required fields
        if not all([department_name, topic, venue, coordinator, summary, outcome]):
            st.error("Please fill in all required fields")
            return
        
        # Ensure time is valid
        if not start_time or not validate_time(start_time):
            st.error("Please enter a valid time in HH:MM format or select from the dropdown")
            return
        
        # Format time for display in the report (keep 12-hour format for better readability)
        try:
            # If time is in 24-hour format, convert to 12-hour for display
            if re.match(r'^([0-1]?[0-9]|2[0-3]):([0-5][0-9])$', start_time):
                time_obj = datetime.strptime(start_time, '%H:%M')
                display_time = time_obj.strftime('%I:%M %p').lstrip('0').replace(' 0', ' ')
            else:
                display_time = start_time
        except:
            display_time = start_time  # Fallback to original input if parsing fails
        
        inputs = {
            'event_type': event_type,
            'department_name': department_name,
            'topic': topic,
            'expert': expert if expert else '',
            'venue': venue,
            'date': date.strftime('%d-%m-%Y'),
            'start_time': display_time,
            'coordinator': coordinator,
            'num_participants': num_participants,
            'summary': summary,
            'outcome': outcome,
            'faculty_in_charge': faculty_in_charge,
            'hod_name': hod_name
        }
        
        files = {
            'invite_image': invite_image,
            'action_photos': action_photos,
            'attendance_photos': attendance_photos,
            'analysis_photos': analysis_photos
        }
        
        try:
            file_data = create_report(inputs, files)
            
            st.download_button(
                label="Download Report",
                data=file_data,
                file_name=f"{inputs['event_type'].lower()}_1_{inputs['department_name']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("Report generated successfully!")
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            st.warning("Please check the inputs and try again.")

if __name__ == "__main__":
    main()
